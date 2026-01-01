#!/usr/bin/env python3
"""
Evaluate a Hugging Face language model (GPT-2, Flan-T5, etc.)
and generate a timestamped run folder with Accuracy, Perplexity, Loss, BLEU, and ROUGE plots
plus a DOCX report.

Terminology clarification:
- "Evaluation Steps (samples)" = number of dataset samples evaluated (no training updates)
- "Training Steps (updates)" = number of optimizer updates during training (used in training-mode runs)

Usage examples:
  python3 test_evaluation.py --model gpt2 --max_steps 10
  python3 test_evaluation.py --model google/flan-t5-small --max_steps 10
"""

import argparse, math, os, datetime
import numpy as np
import torch
import matplotlib.pyplot as plt
from datasets import load_dataset
from transformers import (
    AutoTokenizer,
    AutoModelForCausalLM,
    AutoModelForSeq2SeqLM,
)
import evaluate
from docx import Document
from docx.shared import Inches


# ---------------------------------------------------------------------
# Evaluation core
# ---------------------------------------------------------------------

def evaluate_model(model_name: str, max_steps: int = 10):
    print(f"\n🚀 Evaluating model: {model_name}")
    device = "cuda" if torch.cuda.is_available() else "cpu"

    # --- Load tokenizer ---
    tokenizer = AutoTokenizer.from_pretrained(model_name, use_auth_token=os.getenv("HF_TOKEN"))
    if tokenizer.pad_token is None:
        tokenizer.pad_token = tokenizer.eos_token

    # --- Load model (detect type) ---
    if any(k in model_name.lower() for k in ["t5", "bart", "pegasus", "mt5"]):
        model = AutoModelForSeq2SeqLM.from_pretrained(model_name, use_auth_token=os.getenv("HF_TOKEN")).to(device)
        model_type = "seq2seq"
    else:
        model = AutoModelForCausalLM.from_pretrained(model_name, use_auth_token=os.getenv("HF_TOKEN")).to(device)
        model_type = "causal"
        model.config.pad_token_id = tokenizer.eos_token_id

    # --- Dataset ---
    ds = load_dataset("wikitext", "wikitext-103-raw-v1", split="test[:1%]")
    print(f"Dataset loaded: {len(ds)} samples  |  model_type={model_type}")
    print(f"📊 Evaluation Steps (samples): {max_steps}")

    # --- Optional text-similarity metrics for richer reporting ---
    bleu_metric = evaluate.load("bleu")
    rouge_metric = evaluate.load("rouge")

    preds, refs = [], []
    loss_vals, ppl_vals, acc_vals, bleu_vals, rouge_vals = [], [], [], [], []

    for i, sample in enumerate(ds):
        if i >= max_steps:
            break
        text = sample["text"].strip()
        if not text:
            continue

        # Tokenize input
        inputs = tokenizer(
            text,
            return_tensors="pt",
            truncation=True,
            max_length=128,   # align with SDK defaults
            padding=True,
        ).to(device)

        # --- Compute loss / perplexity ---
        with torch.no_grad():
            outputs = model(
                input_ids=inputs["input_ids"],
                attention_mask=inputs["attention_mask"],
                labels=inputs["input_ids"]
            )
            loss_val = float(outputs.loss.item())
            ppl_val = math.exp(loss_val)
            loss_vals.append(loss_val)
            ppl_vals.append(ppl_val)

        # --- Generate prediction text ---
        with torch.no_grad():
            gen = model.generate(
                inputs["input_ids"],
                attention_mask=inputs["attention_mask"],
                max_new_tokens=32,
                pad_token_id=tokenizer.eos_token_id
            )
        pred_text = tokenizer.decode(gen[0], skip_special_tokens=True)

        preds.append(pred_text)
        refs.append(text)

        # --- Token-level overlap accuracy (SDK-compatible) ---
        ref_tokens = tokenizer.tokenize(text)
        pred_tokens = tokenizer.tokenize(pred_text)
        overlap = len(set(ref_tokens) & set(pred_tokens))
        token_acc = overlap / max(len(ref_tokens), 1)
        acc_vals.append(token_acc)

        # --- Incremental BLEU & ROUGE (normalized per step for plotting) ---
        try:
            bleu_partial = bleu_metric.compute(predictions=preds, references=[[r] for r in refs])["bleu"]
            rouge_partial = rouge_metric.compute(predictions=preds, references=refs).get("rougeL", 0.0)
        except Exception:
            bleu_partial, rouge_partial = 0.0, 0.0
        bleu_vals.append(bleu_partial)
        rouge_vals.append(rouge_partial)

        print(f"[Eval Step {i+1}/{max_steps}] acc={token_acc:.4f} loss={loss_val:.4f} "
              f"ppl={ppl_val:.2f} bleu={bleu_partial:.4f} rougeL={rouge_partial:.4f}")

    # --- Aggregate metrics ---
    acc_mean = float(np.mean(acc_vals)) if acc_vals else 0.0
    loss_mean = float(np.mean(loss_vals)) if loss_vals else 0.0
    ppl_mean = float(np.mean(ppl_vals)) if ppl_vals else 0.0
    bleu_final = bleu_vals[-1] if bleu_vals else 0.0
    rouge_final = rouge_vals[-1] if rouge_vals else 0.0

    print(f"\n✅ Final Evaluation Metrics (over {len(acc_vals)} Evaluation Steps / samples):")
    print(f"Token Accuracy: {acc_mean:.4f}")
    print(f"Loss:           {loss_mean:.4f}")
    print(f"Perplexity:     {ppl_mean:.2f}")
    print(f"BLEU:           {bleu_final:.4f}")
    print(f"ROUGE-L:        {rouge_final:.4f}")

    # Return arrays for plotting and docx
    return (
        np.array(acc_vals),
        np.array(ppl_vals),
        np.array(loss_vals),
        np.array(bleu_vals),
        np.array(rouge_vals),
        model_type,
        max_steps
    )


# ---------------------------------------------------------------------
# Plot helpers
# ---------------------------------------------------------------------

def plot_single_metric(values, label, color, model_name, outdir):
    steps = np.arange(1, len(values) + 1)
    plt.figure(figsize=(6, 4))
    plt.plot(steps, values, label=label, color=color, linewidth=2)
    plt.xlabel("Evaluation Step (sample index)")
    plt.ylabel(label)
    plt.title(f"{label} vs Evaluation Steps — {model_name}")
    plt.grid(alpha=0.3)
    plt.tight_layout()
    out_path = os.path.join(outdir, f"{label.lower()}_{model_name.replace('/', '_')}.png")
    plt.savefig(out_path)
    plt.close()
    print(f"📈 {label} plot saved → {out_path}")
    return out_path


# ---------------------------------------------------------------------
# DOCX Report builder
# ---------------------------------------------------------------------

def build_docx(model_name, model_type, acc_values, ppl_values, loss_values, bleu_values, rouge_values, steps, paths, outdir):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    outfile = os.path.join(outdir, f"evaluation_report_{timestamp}.docx")

    doc = Document()
    doc.add_heading("Language Model Evaluation Report", 0)
    doc.add_paragraph(f"Model: {model_name}")
    doc.add_paragraph(f"Model type: {model_type}")
    doc.add_paragraph(f"Evaluation Steps (samples): {steps}")
    doc.add_paragraph(f"Training Steps (updates): 0  (frozen model; no optimizer updates)")

    # --- Final aggregates ---
    doc.add_heading("Final Metrics", level=1)
    doc.add_paragraph(f"Mean Token Accuracy: {np.mean(acc_values):.4f}")
    doc.add_paragraph(f"Mean Perplexity: {np.mean(ppl_values):.2f}")
    doc.add_paragraph(f"Mean Loss: {np.mean(loss_values):.4f}")
    doc.add_paragraph(f"BLEU: {bleu_values[-1]:.4f}")
    doc.add_paragraph(f"ROUGE-L: {rouge_values[-1]:.4f}")

    # --- Stepwise table ---
    doc.add_heading("Stepwise Metrics", level=1)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Evaluation Step"
    hdr[1].text = "Accuracy"
    hdr[2].text = "Perplexity"
    hdr[3].text = "Loss"
    hdr[4].text = "BLEU"
    hdr[5].text = "ROUGE-L"

    for i, (a, p, l, b, r) in enumerate(zip(acc_values, ppl_values, loss_values, bleu_values, rouge_values), 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = f"{a:.4f}"
        row[2].text = f"{p:.2f}"
        row[3].text = f"{l:.4f}"
        row[4].text = f"{b:.4f}"
        row[5].text = f"{r:.4f}"

    # --- Plots ---
    doc.add_heading("Metric Curves", level=1)
    for label, path in paths.items():
        if path and os.path.exists(path):
            doc.add_paragraph(label)
            doc.add_picture(path, width=Inches(5.5))

    doc.save(outfile)
    print(f"📄 DOCX report saved → {outfile}")
    return outfile


# ---------------------------------------------------------------------
# Main entry
# ---------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--model", type=str, default="gpt2",
                        help="Model name or HF repo ID (e.g., gpt2, google/flan-t5-small)")
    parser.add_argument("--max_steps", type=int, default=10)
    args = parser.parse_args()

    # --- Create timestamped run directory ---
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    base_dir = os.path.join(os.getcwd(), "runs")
    run_dir = os.path.join(base_dir, f"run_{timestamp}")
    os.makedirs(run_dir, exist_ok=True)
    print(f"📁 Created run directory → {run_dir}")

    # --- Run evaluation ---
    acc_values, ppl_values, loss_values, bleu_values, rouge_values, model_type, steps = evaluate_model(args.model, args.max_steps)

    # --- Plots (now include BLEU + ROUGE) ---
    acc_plot = plot_single_metric(acc_values, "Accuracy", "tab:green", args.model, run_dir)
    ppl_plot = plot_single_metric(ppl_values, "Perplexity", "tab:red", args.model, run_dir)
    loss_plot = plot_single_metric(loss_values, "Loss", "tab:orange", args.model, run_dir)
    bleu_plot = plot_single_metric(bleu_values, "BLEU", "tab:blue", args.model, run_dir)
    rouge_plot = plot_single_metric(rouge_values, "ROUGE-L", "tab:purple", args.model, run_dir)

    plots = {
        "Accuracy": acc_plot,
        "Perplexity": ppl_plot,
        "Loss": loss_plot,
        "BLEU": bleu_plot,
        "ROUGE-L": rouge_plot,
    }

    # --- Build DOCX report ---
    build_docx(args.model, model_type, acc_values, ppl_values, loss_values, bleu_values, rouge_values, steps, plots, run_dir)
